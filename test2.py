# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT

document = Document()
section = document.sections[-1]
section.orientation = WD_ORIENT.LANDSCAPE

p = document.add_paragraph(u'''РАСПИСАНИЕ ЗАНЯТИЙ
лаборатории (информационного обеспечения населения и технологий информационной поддержки РСЧС) научно –
исследовательского отдела (по проблемам гражданской обороны и чрезвычайных ситуаций) научно исследовательского центра с 1 декабря по 7 декабря 2014 г.
''')

TitleTable = [u'Дата, дни недели',	u'Подразделение, время проведения',
            u'Предметы обучения', u'Номера тем, занятий, их содержание, отрабатываемые нормативы',
            u'Место проведения', u'Кто проводит', u'Руководства, пособия и материальное обеспечение',
            u'Отметка о проведении']


table = document.add_table(rows=1, cols=len(TitleTable))
table.style = 'TableGrid'
hdr_cells = table.rows[0].cells
for i, a in enumerate(TitleTable):
    hdr_cells[i].text = a

#hdr_cells[0].text = 'Qty'
#hdr_cells[1].text = 'Id'
#hdr_cells[2].text = 'Desc'
#row_cells = table.add_row(TitleTable)

row_cells = table.add_row().cells
row_cells[0].text = str('fef')
row_cells[1].text = str('fef')
row_cells[2].text = 'fef'
row_cells = table.add_row().cells

document.add_page_break()

document.save('demo.docx')